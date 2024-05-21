import streamlit as st
from jinja2 import Environment, PackageLoader, select_autoescape, FileSystemLoader
import os
from streamlit.components.v1 import iframe
import matplotlib.pyplot as plt
import numpy as np
from PIL import Image
import io
from datetime import datetime, date

st.title("PUMP PERFORMANCE REPORT GENERATOR")
form = st.form("template_form")

v1= form.date_input("DATE OF INSPECTION")
v2 = form.text_input("ENTER PROJECT NAME")
v3 = form.text_input("ENTER PUMP NAME TAG",value="PUMP NAME TAG")
v4 = form.text_input("ENTER Pump Make")
v5 = form.text_input("ENTER Pump MODEL NO")
v6 = form.text_input("ENTER Pump Sr. No.")

v8 = form.number_input(
    label="ENTER RATED PUMP FLOW RATE IN M^3/H",
    min_value=0.0,  # minimum value allowed
    max_value=800.0,  # maximum value allowed
    value=300.0,  # default value
    step=0.1,  # step size
    format="%.2f"  # number format
)

v7 = form.number_input(
    label="ENTER RATED PUMP HEAD IN M",
    min_value=0.0,  # minimum value allowed
    max_value=100.0,  # maximum value allowed
    value=14.0,  # default value
    step=0.1,  # step size
    format="%.2f"  # number format
)




v10 =round( form.number_input(
    label="Pump Suction pressure in Kg/cm2",
    min_value=0.0,  # minimum value allowed
    max_value=20.0,  # maximum value allowed
    value=2.5,  # default value
    step=0.1,  # step size
    format="%.2f"  # number format
),2)



v11 =round( form.number_input(
    label="Pump Discharge pressure in Kg/cm2",
    min_value=0.0,  # minimum value allowed
    max_value=20.0,  # maximum value allowed
    value=4.0,  # default value
    step=0.1,  # step size
    format="%.2f"  # number format
),2)


v12=round((v11-v10)*10,2)

# cmh=500
# rph=50
# my=40

cmh=v8
rph=v7
my=v12

b=1.32*rph
a=(rph-b)/(cmh**2)

# Rated head and flow BEP point

plt.scatter(cmh, rph, color='green', label='Rated Selection Point')

#HEAD FLOW CURVE

x1=np.linspace(0, 1000, 100)
y1=(a*x1**2)+b

y=[]
for i in y1:
    if i>=0:
        y.append(i)

x=x1[0:len(y)]

plt.plot(x, y)

mflow=np.sqrt((my-b)/a)

v13 = round(form.number_input(
    label="Measured flowrate in m^3/h",
    min_value=0.0,  # minimum value allowed
    max_value=cmh+cmh,  # maximum value allowed
    value=mflow,  # default value
    step=0.1,  # step size
    format="%.2f"  # number format
),2)


#ACTUAL SYSTEM CURVE


sx=np.linspace(0, cmh,100)
sy=rph*((sx/cmh)**2)
# print(y)
plt.plot(sx, sy, label='Actual System Curve')


#OPERATING  SYSTEM CURVE

tsx=np.linspace(0, mflow, 100)
tsy=(my)*((tsx/mflow)**2)
# print(y)
plt.plot(tsx, tsy, label='Operating System Curve')


x_values = [0, mflow]  # x-values for the line
y_value = my        # y-value for the line

# Define the point
point_x = mflow
point_y = my


#HORIZOTAL THORTTLE  LINE

# Plot the horizontal THROTTLE  line
plt.plot(x_values, [y_value, y_value], color='blue', label='Operating Head')

# Plot the OPERARING point
plt.scatter(point_x, point_y, color='red', label='Operating Point')


# VERTICAL flow LINE

plt.plot([mflow,mflow],[my,0],color='red', label='Operating flow')





# Add labels and title
plt.xlabel('Flow Rate (m³/hr)')
plt.ylabel('Head (m)')
plt.title('Pump Head-Flow Curve')


plt.annotate("BEP",
                  xy=(cmh, rph), xycoords='data',
                  xytext=(cmh+(cmh*10/100),rph+(rph*10/100)), textcoords='data',
                  size=10, va="center", ha="center",
                  bbox=dict(boxstyle="round4", fc="w"),
                  arrowprops=dict(arrowstyle="-|>",
                                  connectionstyle="arc3,rad=-0.2",
                                  fc="w"))

# Add a legend
plt.legend(loc='upper right',fontsize=8)
plt.xlim(0,cmh+cmh)
# Display the plot
plt.grid(True)
# plt.figure(figsize=(5, 5))
# plt.savefig("trial.png")
plt.savefig('my_plot.png', dpi=300)  # Higher resolution
plt.show()


# Calculation

v9=round(((v8/3600)*1000*9.81*v7)/1000,2)
v14=round((((v13/3600)*1000*9.81*v12)/1000),2)


v15=form.number_input("RATED_MOTOR_POWER_IN_KW",value=15)

v18=form.slider("RATED_MOTOR_EFFICIENCY", 50.0, 100.0,90.0,step=0.5)
v16=form.slider("RATED_MOTOR_POWER_FACTOR", 0.7, 1.0,0.89,step=0.05)
v17=round(form.number_input("SELECTED_RATED_CURRENT",value=((1000*v15)/(1.732*415*0.9))),2)

v19=form.number_input("MEASURED_VOLTAGE_R_AND_Y",value=415)
v20=form.number_input("MEASURED_VOLTAGE_Y_AND_B",value=415)
v21=form.number_input("MEASURED_VOLTAGE_B_AND_R",value=415)
v22=round(form.number_input("MEASURED_CURRENT_R",value=v17-3),2)
v23=round(form.number_input("MEASURED_CURRENT_Y",value=v17-3),2)
v24=round(form.number_input("MEASURED_CURRENT_B",value=v17-3),2)


v25=round((1.732*((v19+v20+v21)/3)*((v22+v23+v24)/3)*v16)/1000,2)

v27=round((v25*v18/100),2)
v28=round((v14/v27)*100,2)
v29=round((1.732*v17*415*v16)/1000,2)

ref=round((v9/v15)*100,2)
v26=round(form.slider("SELECTED_PUMP_EFFICIENCY", 30.0, 100.0,ref,step=0.5),2)    

import docx
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from io import BytesIO

document = Document()
document.add_heading('PUMP PERFORMANCE CURVE', 0)
document.add_picture('my_plot.png')  
    
##SIGNOFF

document.add_heading('SIGNOFF......', level=1)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'ATTENDEES'
hdr_cells[1].text = 'PARTIES'
hdr_cells[2].text = 'SIGNATURE'   



submit = form.form_submit_button("CALCULATE & GENERATE REPORTS")

env = Environment(loader=FileSystemLoader("."), autoescape=select_autoescape())
template = env.get_template("rep.html")
data=[v1,v2,v3,v4,v5,v6,v7,v8,v9,v10,v11,v12,v13,v14,v15,v16,v17]

im = Image.open("my_plot.png")
image_bytes = io.BytesIO()

if submit:
    st.image("my_plot.png")
    html = template.render(V1=v1,
V2=v2,
V3=v3,
V4=v4,
V5=v5,
V6=v6,
V7=v7,
V8=v8,
V9=v9,
V10=v10,
V11=v11,
V12=v12,
V13=v13,
V14=v14,
V15=v15,
V16=v16,
V17=v17,
V18=v18,
V19=v19,
V20=v20,
V21=v21,
V22=v22,
V23=v23,
V24=v24,
V25=v25,
V26=v26,
V27=v27,
V28=v28,
V29=v29,
)

    with open("biju.html","w",encoding="utf-8") as f:
        f.write(html)
    f.close()
    # Save the pdf with name
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    
    #DATABASE
    import pymongo
    client = pymongo.MongoClient("mongodb+srv://biju:biju123@cluster0.tu5mpny.mongodb.net/?retryWrites=true&w=majority&appName=Cluster0")
    db = client.d27
    records=db.d27coll

    # for x in records.find({}):
    # print(x)


    
    new={'Date_of_Testing':datetime.combine(v1, datetime.min.time()),
    'Project_name':v2,
    'Pump_Name_Tag':v3,
    'Pump Make':v4,
    'Pump Model':v5,
    'Pump Sr. No.':v6,
    'HYDRAULIC PARAMETERS': {"RATED":{'Pump Suction pressure in Kg/cm2': "-",
                'Pump Discharge pressure in Kg/cm2': "-",
                'Pump Head in m': v7,
                'Flow rate in m^3/h': v8,
                'Hydraulic output Power': v9
            },
                "MEASURED":{'Pump Suction pressure in Kg/cm2':v10,
                'Pump Discharge pressure in Kg/cm2':v11,
                'Pump Head in m': v12,
                'Flow rate in m^3/h': v13,
                'Hydraulic output Power': v14
            }
    },
    'INPUT POWER PARAMETERS': {"RATED":{'Rated_Motor_Power_in _KW':v15,
                'Motor_Power_Factor': v16,
                'Rated_Motor_Current_in _A': v17,
                'Rated_Motor_Efficieny': v18,
                
            },
                "MEASURED":{'Measured R_Y_Voltage':v19,
                'Measured Y_B_Voltage':v20,
                'Measured B_R_Voltage':v21,
                'Measured R_Current': v22,
                'Measured Y_Current': v23,
                'Measured B_Current': v24
            }
    },
        'EFFIECIENCY PARAMETERS': {"RATED":{'INPUT ELECTRICAL POWER':v29,
                'PUMP_HYDRAULIC _EFFICIENCY': v26,
            },
                "MEASURED":{'INPUT ELECTRICAL POWER':v25,
                'MOTOR SHAFT OUTPUT POWER': v27,
                'PUMP_HYDRAULIC _EFFICIENCY': v28,
    }
    },'curve':image_bytes.getvalue()}

    records.insert_one(new)
    # for x in records.find({}):
    #     print(x)

    
    document.save("curve.docx")
    
    
    
    # st.download_button("⬇️ Download Calculated Report",data=html,file_name=v3+".html")
    # st.download_button(label="⬇️ Download Performance Curve Report",data=buffer, file_name='Pump Curve_REPORT.docx',
    #                    mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    
    import streamlit as st
    import zipfile
    import io

    def create_zip(files):
        # Create an in-memory bytes buffer
        buffer = io.BytesIO()
        
        # Create a zip file in the buffer
        with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for file_name, data in files.items():
                # Ensure data is in bytes format
                if isinstance(data, str):
                    data = data.encode('utf-8')
                zipf.writestr(file_name, data)
        
        # Move the buffer cursor to the beginning
        buffer.seek(0)
        return buffer

    # Example data to be included in the zip file
    file1_content = io.BytesIO(open("biju.html", "rb").read()).getvalue()
    file2_content = io.BytesIO(open("curve.docx", "rb").read()).getvalue()

    # Dictionary with filenames as keys and file content as values
    files_to_zip = {
        v3+".html": file1_content,
        "'Pump Curve_REPORT.docx'": file2_content
    }

    # Create the zip file
    zip_buffer = create_zip(files_to_zip)


    # Provide a download button for the zip file
    st.download_button(
        label="Download All Reports Files as Zip file for Printing",
        data=zip_buffer,
        file_name=v2+"-"+v3+"-"+v6+".zip",
        mime="application/zip"
    )
